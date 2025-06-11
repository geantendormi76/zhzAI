import os
from docx import Document
from docx.shared import Inches # 用于处理图片（如果需要提取图片信息）
import logging

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

SAMPLE_DOCS_DIR = "sample_docs"

def analyze_docx_with_python_docx(file_path):
    logger.info(f"\n--- Analyzing DOCX: {file_path} ---")
    try:
        doc = Document(file_path)
        
        # 1. 提取所有段落文本及样式信息
        logger.info("\n[Paragraphs Text and Styles]")
        for i, para in enumerate(doc.paragraphs):
            style_name = para.style.name if para.style else "No Style"
            # 尝试获取更具体的列表信息（如果段落属于列表）
            list_info = ""
            if para.style.name.startswith('List Paragraph'): # Word 内置列表段落样式
                try:
                    # 这部分比较复杂，简单的 para.paragraph_format.left_indent 可能不足以判断级别
                    # python-docx 对于列表级别的直接提取能力有限
                    # 我们可以检查是否存在编号相关的 XML 元素，但这超出了简单PoC
                    # 暂时仅通过样式名提示
                    list_info = " (Potential List Item)"
                except Exception:
                    pass # 忽略获取列表信息的错误
            logger.info(f"  Para {i+1} (Style: '{style_name}'){list_info}: \"{para.text[:150]}{'...' if len(para.text) > 150 else ''}\"")


        # 2. 提取表格内容
        if doc.tables:
            logger.info(f"\n[Tables Content (Found {len(doc.tables)} table(s))]")
            for i, table in enumerate(doc.tables):
                logger.info(f"  Table {i+1}:")
                for row_idx, row in enumerate(table.rows):
                    row_data = [cell.text.strip() for cell in row.cells]
                    logger.info(f"    Row {row_idx+1}: {row_data}")
        else:
            logger.info("\n[Tables Content]: No tables found in this document.")

        # 3. 提取页眉和页脚 (修正了索引获取方式)
        logger.info("\n[Headers and Footers (Basic)]")
        for section_idx, section in enumerate(doc.sections): # 使用 enumerate 获取索引
            if section.header:
                header_text = " ".join([p.text for p in section.header.paragraphs if p.text.strip()])
                if header_text:
                     logger.info(f"  Header (Section {section_idx+1}): \"{header_text[:100]}{'...' if len(header_text) > 100 else ''}\"")
                else:
                    logger.info(f"  Header (Section {section_idx+1}): (empty or not directly accessible text)")
            else:
                logger.info(f"  Header (Section {section_idx+1}): Not found")

            if section.footer:
                footer_text = " ".join([p.text for p in section.footer.paragraphs if p.text.strip()])
                if footer_text:
                    logger.info(f"  Footer (Section {section_idx+1}): \"{footer_text[:100]}{'...' if len(footer_text) > 100 else ''}\"")
                else:
                    logger.info(f"  Footer (Section {section_idx+1}): (empty or not directly accessible text)")
            else:
                logger.info(f"  Footer (Section {section_idx+1}): Not found")
        
        # 4. 提取图片信息 (python-docx 主要关注内联图片)
        logger.info("\n[Inline Shapes (Potential Images)]")
        inline_shapes_count = 0
        
        # 1. 检查文档主体中的内联形状
        if doc.inline_shapes:
            for shape in doc.inline_shapes:
                inline_shapes_count += 1
                logger.info(f"  Found inline shape in document body: Type={shape.type}, Width={shape.width}, Height={shape.height}")
        
        # 2. 检查页眉和页脚中的段落内是否存在图片
        # python-docx 对直接访问页眉/页脚内复杂对象（如图形）的支持有限
        # 更可靠的方法是迭代段落和runs
        for section_idx, section in enumerate(doc.sections):
            if section.header:
                for paragraph in section.header.paragraphs:
                    for run in paragraph.runs:
                        # 'drawing' 通常代表嵌入的图片或其他绘图对象
                        # run.element.xpath('.//wp:docPr') 可以用来获取图片的描述或标题 (如果设置了)
                        # 但这里我们只做简单的存在性检测
                        if run.element.tag.endswith('drawing'): # 检查run是否包含绘图对象
                            inline_shapes_count +=1
                            logger.info(f"  Found drawing element (potential image) in header (Section {section_idx+1}), paragraph: '{paragraph.text[:30]}...'")
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    for run in paragraph.runs:
                        if run.element.tag.endswith('drawing'):
                            inline_shapes_count +=1
                            logger.info(f"  Found drawing element (potential image) in footer (Section {section_idx+1}), paragraph: '{paragraph.text[:30]}...'")
                            
        if inline_shapes_count == 0:
             logger.info("  No inline shapes or drawing elements detected in document body, headers, or footers.")

    except Exception as e:
        logger.error(f"Error processing file {file_path}: {e}", exc_info=True)

if __name__ == "__main__":
    if not os.path.isdir(SAMPLE_DOCS_DIR):
        logger.error(f"Sample documents directory '{SAMPLE_DOCS_DIR}' not found. Please create it and add .docx files.")
    else:
        # 获取目录下所有.docx文件并排序，确保处理顺序一致
        docx_files = sorted([
        f for f in os.listdir(SAMPLE_DOCS_DIR) 
        if f.lower().endswith(".docx") and not f.startswith("~$")
    ])
        if not docx_files:
            logger.warning(f"No .docx files found in '{SAMPLE_DOCS_DIR}'.")
        else:
            for filename in docx_files:
                file_path = os.path.join(SAMPLE_DOCS_DIR, filename)
                analyze_docx_with_python_docx(file_path)